from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os
import logging
from datetime import datetime


class WordExporter:
    FONT_NAME = 'MS Mincho'
    FONT_NAME_FALLBACK = 'ＭＳ 明朝'

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def _set_run_font(self, run, font_name=None, font_size=None, bold=False, underline=False, color=None):
        fn = font_name or self.FONT_NAME
        run.font.name = fn
        run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
        if font_size:
            run.font.size = font_size
        if bold:
            run.font.bold = True
        if underline:
            run.font.underline = True
        if color:
            run.font.color.rgb = color

    def export_single(self, essay, save_path):
        try:
            doc = Document()
            self._setup_document(doc)

            # 标题：MC Mincho 三号(16pt) 粗体 居中
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(essay.get('title', '无标题'))
            self._set_run_font(title_run, font_size=Pt(16), bold=True)
            title_para.paragraph_format.space_after = Pt(0)
            title_para.paragraph_format.space_before = Pt(0)

            # 作者：MC Mincho 小五(9pt) 下划线 居中
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_text = essay.get('author', '未知')
            if essay.get('school'):
                author_text += f" （{essay['school']}）"
            author_run = author_para.add_run(author_text)
            self._set_run_font(author_run, font_size=Pt(9), underline=True)
            author_para.paragraph_format.space_after = Pt(0)
            author_para.paragraph_format.space_before = Pt(0)

            # 正文：MC Mincho 小四(12pt)
            body_para = doc.add_paragraph()
            body_para.paragraph_format.line_spacing = 1.8
            body_para.paragraph_format.space_after = Pt(0)
            body_para.paragraph_format.space_before = Pt(0)
            body_text = essay.get('body', '')
            # 按段落分割，保持原文段落结构
            body_parts = body_text.split('\n\n')
            for idx, part in enumerate(body_parts):
                if not part.strip():
                    continue
                if idx > 0:
                    body_para.add_run('\n')
                run = body_para.add_run(part.strip())
                self._set_run_font(run, font_size=Pt(12))

            # 网址：MC Mincho 小五(9pt) 灰色 右对齐
            source_para = doc.add_paragraph()
            source_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            source_run = source_para.add_run(essay.get('source', ''))
            self._set_run_font(source_run, font_size=Pt(9), color=RGBColor(0x80, 0x80, 0x80))
            source_para.paragraph_format.space_after = Pt(0)
            source_para.paragraph_format.space_before = Pt(0)

            doc.save(save_path)
            self.logger.info(f'Exported: {save_path}')
            return True
        except Exception as e:
            self.logger.error(f'Export error: {save_path} - {str(e)}')
            return False

    def export_batch(self, essays, save_path, date_from=None, date_to=None):
        try:
            doc = Document()
            self._setup_document(doc)

            for i, essay in enumerate(essays, 1):
                # 标题
                title_para = doc.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_para.add_run(f"{i}. {essay.get('title', '无标题')}")
                self._set_run_font(title_run, font_size=Pt(16), bold=True)
                title_para.paragraph_format.space_after = Pt(0)
                title_para.paragraph_format.space_before = Pt(0)

                # 作者
                author_para = doc.add_paragraph()
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                author_text = essay.get('author', '未知')
                if essay.get('school'):
                    author_text += f" （{essay['school']}）"
                author_run = author_para.add_run(author_text)
                self._set_run_font(author_run, font_size=Pt(9), underline=True)
                author_para.paragraph_format.space_after = Pt(0)
                author_para.paragraph_format.space_before = Pt(0)

                # 正文
                body_para = doc.add_paragraph()
                body_para.paragraph_format.line_spacing = 1.8
                body_para.paragraph_format.space_after = Pt(0)
                body_para.paragraph_format.space_before = Pt(0)
                body_text = essay.get('body', '')
                body_parts = body_text.split('\n\n')
                for idx, part in enumerate(body_parts):
                    if not part.strip():
                        continue
                    if idx > 0:
                        body_para.add_run('\n')
                    run = body_para.add_run(part.strip())
                    self._set_run_font(run, font_size=Pt(12))

                # 网址
                source_para = doc.add_paragraph()
                source_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                source_run = source_para.add_run(essay.get('source', ''))
                self._set_run_font(source_run, font_size=Pt(9), color=RGBColor(0x80, 0x80, 0x80))
                source_para.paragraph_format.space_after = Pt(0)
                source_para.paragraph_format.space_before = Pt(0)

                if i < len(essays):
                    doc.add_page_break()

            doc.save(save_path)
            self.logger.info(f'Batch exported {len(essays)} essays: {save_path}')
            return True
        except Exception as e:
            self.logger.error(f'Batch export error: {save_path} - {str(e)}')
            return False

    def _setup_document(self, doc):
        # 设置默认段落样式，消除段间距
        style = doc.styles['Normal']
        font = style.font
        font.name = self.FONT_NAME
        font.size = Pt(12)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_NAME)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)

        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
